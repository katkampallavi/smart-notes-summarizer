import os
import re


def extract_text_from_docx(filepath):
    """
    Extract text content from a .docx file using python-docx.
    Returns extracted text as a string.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"DOCX file not found: {filepath}")

    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(' | '.join(row_texts))

        text = '\n\n'.join(paragraphs)

        if not text.strip():
            raise ValueError("The DOCX file appears to be empty or contains no extractable text.")

        return _clean_extracted_text(text)

    except ImportError:
        raise ImportError(
            "python-docx is required to read DOCX files. "
            "Install it with: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX file: {str(e)}")


def _clean_extracted_text(text):
    """Clean and normalize extracted text from DOCX."""
    # Remove non-printable characters except common whitespace
    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
    # Normalize spaces
    text = re.sub(r'[ \t]+', ' ', text)
    # Normalize newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.strip()
    return text